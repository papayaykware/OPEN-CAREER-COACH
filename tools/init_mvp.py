#!/usr/bin/env python
# init_mvp.py
"""
Generador limpio del MVP de OPEN CAREER COACH.

Uso:
    python init_mvp.py --path ./open-career-coach
"""

import argparse
import os
from pathlib import Path
from textwrap import dedent

# -------------------------
# Configuración del layout
# -------------------------

DIRECTORIES = [
    "src",
    "src/utils",
    "src/cv_parser",
    "src/job_parser",
    "src/matching",
    "src/ui",
    "data",
    "data/sample_cvs",
    "data/sample_jobs",
    "tests",
]

INIT_PACKAGES = [
    "src",
    "src/utils",
    "src/cv_parser",
    "src/job_parser",
    "src/matching",
    "src/ui",
]

# -------------------------
# Plantillas mínimas
# -------------------------

REQUIREMENTS_TXT = dedent(
    """\
    # OPEN CAREER COACH - MVP Dependencies

    # Core
    gradio==4.44.0
    sentence-transformers==3.0.1
    PyPDF2==3.0.1
    python-docx==1.1.2
    chromadb==0.5.5
    numpy==1.26.4

    # NLP
    spacy==3.7.5

    # Utilities
    python-dotenv==1.0.1
    requests==2.32.3

    # Testing
    pytest==8.3.2
    """
)

CONFIG_PY = dedent(
    """\
    \"\"\"Configuración central del sistema MVP.\"\"\"

    from dataclasses import dataclass


    @dataclass
    class Config:
        \"\"\"Configuración global.\"\"\"

        # Modelos
        EMBEDDING_MODEL: str = "sentence-transformers/all-MiniLM-L6-v2"

        # Matching
        SIMILARITY_THRESHOLD: float = 0.50

        # Paths
        DATA_DIR: str = "./data"
        SAMPLE_CVS_DIR: str = "./data/sample_cvs"
        SAMPLE_JOBS_DIR: str = "./data/sample_jobs"


    config = Config()
    """
)

FILE_LOADER_PY = dedent(
    """\
    \"\"\"Carga y extrae texto de archivos PDF, DOCX y TXT.\"\"\"

    from pathlib import Path
    from typing import Dict


    def load_text_from_file(filepath: str) -> Dict[str, str]:
        \"\"\"Detecta la extensión del archivo y extrae el texto plano.

        Args:
            filepath: Ruta al archivo (PDF, DOCX o TXT)

        Returns:
            Dict con 'text', 'format', 'pages_estimated', 'characters'
        \"\"\"
        path = Path(filepath)

        if not path.exists():
            raise FileNotFoundError(f"Archivo no encontrado: {filepath}")

        suffix = path.suffix.lower()

        if suffix == ".pdf":
            return _extract_pdf(filepath)
        elif suffix == ".docx":
            return _extract_docx(filepath)
        elif suffix == ".txt":
            return _extract_txt(filepath)
        else:
            raise ValueError(f"Formato no soportado: {suffix}. Use PDF, DOCX o TXT.")


    def _extract_pdf(filepath: str) -> Dict[str, str]:
        import PyPDF2

        text_parts = []

        with open(filepath, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            total_pages = len(reader.pages)

            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

        full_text = "\\n\\n".join(text_parts)

        return {
            "text": full_text,
            "format": "PDF",
            "pages_estimated": total_pages,
            "characters": len(full_text),
        }


    def _extract_docx(filepath: str) -> Dict[str, str]:
        from docx import Document

        doc = Document(filepath)

        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    table_texts.append(row_text)

        full_text = "\\n\\n".join(paragraphs + table_texts)

        return {
            "text": full_text,
            "format": "DOCX",
            "pages_estimated": len(paragraphs) // 40 + 1,
            "characters": len(full_text),
        }


    def _extract_txt(filepath: str) -> Dict[str, str]:
        with open(filepath, "r", encoding="utf-8") as f:
            text = f.read()

        if not text.strip():
            for encoding in ["latin-1", "iso-8859-1", "cp1252"]:
                try:
                    with open(filepath, "r", encoding=encoding) as f:
                        text = f.read()
                    if text.strip():
                        break
                except Exception:
                    continue

        lines = text.splitlines()

        return {
            "text": text,
            "format": "TXT",
            "pages_estimated": len(lines) // 50 + 1,
            "characters": len(text),
        }
    """
)

CV_PIPELINE_PY = dedent(
    """\
    \"\"\"Pipeline de procesamiento de currículums para el MVP.\"\"\"

    import re
    from typing import List, Optional
    from dataclasses import dataclass, field

    from src.utils.file_loader import load_text_from_file


    @dataclass
    class CVData:
        raw_text: str = ""
        skills_technical: List[str] = field(default_factory=list)
        skills_soft: List[str] = field(default_factory=list)
        experience_years: Optional[float] = None
        email: Optional[str] = None
        phone: Optional[str] = None
        linkedin: Optional[str] = None
        file_format: str = ""
        pages: int = 0


    class CVPipeline:
        \"\"\"Pipeline completo de procesamiento de CVs.\"\"\"

        TECH_SKILLS = {
            "python", "java", "javascript", "typescript",
            "sql", "docker", "kubernetes", "aws",
        }

        SOFT_SKILLS = {
            "liderazgo", "leadership", "comunicación", "communication",
            "trabajo en equipo", "teamwork",
        }

        EMAIL_PATTERN = re.compile(r"\\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\\.[A-Z|a-z]{2,}\\b")
        PHONE_PATTERN = re.compile(r"(?:\\+?\\d{1,3}[-.\\s]?)?\\(?\\d{2,4}\\)?[-.\\s]?\\d{2,4}[-.\\s]?\\d{2,4}")
        LINKEDIN_PATTERN = re.compile(r"linkedin\\.com/in/[\\w-]+")

        def __init__(self) -> None:
            self.tech_skills_lower = {s.lower() for s in self.TECH_SKILLS}
            self.soft_skills_lower = {s.lower() for s in self.SOFT_SKILLS}

        def process(self, file_path: str) -> CVData:
            extraction = load_text_from_file(file_path)
            text = extraction["text"]

            cv_data = CVData(
                raw_text=text,
                file_format=extraction["format"],
                pages=extraction.get("pages_estimated", 0),
            )

            cv_data.email = self._extract_email(text)
            cv_data.phone = self._extract_phone(text)
            cv_data.linkedin = self._extract_linkedin(text)

            tech, soft = self._extract_skills(text)
            cv_data.skills_technical = tech
            cv_data.skills_soft = soft

            cv_data.experience_years = self._estimate_experience(text)

            return cv_data

        def _extract_email(self, text: str) -> Optional[str]:
            match = self.EMAIL_PATTERN.search(text)
            return match.group() if match else None

        def _extract_phone(self, text: str) -> Optional[str]:
            match = self.PHONE_PATTERN.search(text)
            return match.group() if match else None

        def _extract_linkedin(self, text: str) -> Optional[str]:
            match = self.LINKEDIN_PATTERN.search(text)
            return match.group() if match else None

        def _extract_skills(self, text: str) -> tuple[list[str], list[str]]:
            text_lower = text.lower()

            tech_found = []
            for skill in self.tech_skills_lower:
                pattern = r"\\b" + re.escape(skill) + r"\\b"
                if re.search(pattern, text_lower):
                    tech_found.append(skill)

            soft_found = []
            for skill in self.soft_skills_lower:
                pattern = r"\\b" + re.escape(skill) + r"\\b"
                if re.search(pattern, text_lower):
                    soft_found.append(skill)

            return sorted(set(tech_found)), sorted(set(soft_found))

        def _estimate_experience(self, text: str) -> Optional[float]:
            text_lower = text.lower()
            patterns = [
                r"(\\d+)\\+?\\s*(?:years?|años?)(?:\\s+of)?(?:\\s+experience)?",
                r"(?:experiencia(?:\\s+de)?\\s+)(\\d+)\\s*(?:años?|years?)",
            ]

            for pattern in patterns:
                matches = re.findall(pattern, text_lower)
                if matches:
                    try:
                        years = [int(m) for m in matches if m.isdigit()]
                        if years:
                            return float(min(years))
                    except Exception:
                        continue

            return None
    """
)

JOB_PIPELINE_PY = dedent(
    """\
    \"\"\"Pipeline de procesamiento de ofertas de empleo para el MVP.\"\"\"

    import re
    from typing import List, Optional, Tuple
    from dataclasses import dataclass, field


    @dataclass
    class JobData:
        raw_text: str = ""
        title: str = ""
        company: str = ""
        required_skills: List[str] = field(default_factory=list)
        soft_skills: List[str] = field(default_factory=list)
        experience_years: Optional[int] = None
        experience_range: Optional[Tuple[int, int]] = None
        remote_policy: str = "unknown"
        employment_type: str = "full-time"


    class JobPipeline:
        TECH_SKILLS = {
            "python", "java", "javascript", "typescript",
            "sql", "docker", "kubernetes", "aws",
        }

        SOFT_SKILLS = {
            "liderazgo", "leadership", "comunicación", "communication",
            "trabajo en equipo", "teamwork",
        }

        EXPERIENCE_PATTERNS = [
            r"(?:minimum\\s+)?(?:(\\d+)\\+?\\s*(?:years?|años?)(?:\\s+of)?(?:\\s+experience)?)",
            r"(?:(\\d+)\\s*(?:-|to|a)\\s*(\\d+)\\s*(?:years?|años?))",
        ]

        def __init__(self) -> None:
            self.tech_skills_lower = {s.lower() for s in self.TECH_SKILLS}
            self.soft_skills_lower = {s.lower() for s in self.SOFT_SKILLS}

        def process(self, text: str) -> JobData:
            job = JobData(raw_text=text)
            clean_text = self._preprocess(text)

            tech, soft = self._extract_skills(clean_text)
            job.required_skills = tech
            job.soft_skills = soft

            job.experience_years, job.experience_range = self._extract_experience(clean_text)

            job.remote_policy = self._extract_remote_policy(clean_text)
            job.employment_type = self._extract_employment_type(clean_text)

            return job

        def _preprocess(self, text: str) -> str:
            text = re.sub(r"<[^>]+>", " ", text)
            text = re.sub(r"\\s+", " ", text)
            return text.strip()

        def _extract_skills(self, text: str) -> tuple[list[str], list[str]]:
            text_lower = text.lower()

            tech_found = []
            for skill in self.tech_skills_lower:
                pattern = r"\\b" + re.escape(skill) + r"\\b"
                if re.search(pattern, text_lower):
                    tech_found.append(skill)

            soft_found = []
            for skill in self.soft_skills_lower:
                pattern = r"\\b" + re.escape(skill) + r"\\b"
                if re.search(pattern, text_lower):
                    soft_found.append(skill)

            return sorted(set(tech_found)), sorted(set(soft_found))

        def _extract_experience(self, text: str) -> Tuple[Optional[int], Optional[Tuple[int, int]]]:
            text_lower = text.lower()

            for pattern in self.EXPERIENCE_PATTERNS:
                matches = re.findall(pattern, text_lower)
                for match in matches:
                    if isinstance(match, tuple):
                        try:
                            min_years = int(match[0])
                            max_years = int(match[1])
                            return None, (min_years, max_years)
                        except Exception:
                            continue
                    else:
                        try:
                            years = int(match)
                            return years, None
                        except Exception:
                            continue

            return None, None

        def _extract_remote_policy(self, text: str) -> str:
            text_lower = text.lower()

            if any(kw in text_lower for kw in ["100% remote", "fully remote", "remoto 100%", "teletrabajo"]):
                return "remote"
            if any(kw in text_lower for kw in ["hybrid", "híbrido", "mixto"]):
                return "hybrid"
            if any(kw in text_lower for kw in ["on-site", "presencial", "oficina"]):
                return "on-site"

            return "unknown"

        def _extract_employment_type(self, text: str) -> str:
            text_lower = text.lower()

            if any(kw in text_lower for kw in ["full-time", "tiempo completo", "jornada completa"]):
                return "full-time"
            if any(kw in text_lower for kw in ["part-time", "tiempo parcial", "media jornada"]):
                return "part-time"
            if any(kw in text_lower for kw in ["contract", "contrato", "freelance", "autónomo"]):
                return "contract"
            if any(kw in text_lower for kw in ["internship", "prácticas", "becario", "intern"]):
                return "internship"

            return "full-time"
    """
)

SIMILARITY_PY = dedent(
    """\
    \"\"\"Motor de matching semántico entre CVs y ofertas de empleo.\"\"\"

    from dataclasses import dataclass, field
    from typing import Dict, List, Optional, Tuple

    import numpy as np
    from sentence_transformers import SentenceTransformer


    @dataclass
    class MatchResult:
        global_score: float
        semantic_similarity: float
        skill_match_score: float
        experience_match_score: float
        matched_skills: List[str] = field(default_factory=list)
        missing_skills: List[str] = field(default_factory=list)
        extra_skills: List[str] = field(default_factory=list)
        recommendations: List[str] = field(default_factory=list)


    class CVJobMatcher:
        \"\"\"Motor de matching entre currículums y ofertas de empleo.\"\"\"

        def __init__(self, model_name: str = "sentence-transformers/all-MiniLM-L6-v2") -> None:
            self.model = SentenceTransformer(model_name)

        def calculate_match(self, cv_data: Dict, job_data: Dict) -> MatchResult:
            semantic_sim = self._calculate_semantic_similarity(
                cv_data.get("raw_text", ""),
                job_data.get("raw_text", ""),
            )

            skill_score, matched, missing, extra = self._calculate_skill_match(
                cv_data.get("skills_technical", []),
                job_data.get("required_skills", []),
            )

            exp_score = self._calculate_experience_match(
                cv_data.get("experience_years"),
                job_data.get("experience_years"),
                job_data.get("experience_range"),
            )

            global_score = 0.40 * semantic_sim + 0.40 * skill_score + 0.20 * exp_score

            recommendations = self._generate_recommendations(missing, cv_data, job_data)

            return MatchResult(
                global_score=round(global_score, 1),
                semantic_similarity=round(semantic_sim, 1),
                skill_match_score=round(skill_score, 1),
                experience_match_score=round(exp_score, 1),
                matched_skills=matched,
                missing_skills=missing,
                extra_skills=extra,
                recommendations=recommendations,
            )

        def _calculate_semantic_similarity(self, cv_text: str, job_text: str) -> float:
            if not cv_text or not job_text:
                return 0.0

            embeddings = self.model.encode([cv_text, job_text])
            cv_emb, job_emb = embeddings[0], embeddings[1]

            num = float(np.dot(cv_emb, job_emb))
            den = float(np.linalg.norm(cv_emb) * np.linalg.norm(job_emb))
            if den == 0:
                return 0.0

            cos_sim = num / den
            return max(0.0, min(1.0, cos_sim)) * 100.0

        def _calculate_skill_match(
            self,
            cv_skills: List[str],
            job_skills: List[str],
        ) -> Tuple[float, List[str], List[str], List[str]]:
            cv_set = {s.lower() for s in cv_skills}
            job_set = {s.lower() for s in job_skills}

            matched = sorted(cv_set & job_set)
            missing = sorted(job_set - cv_set)
            extra = sorted(cv_set - job_set)

            if not job_set:
                score = 100.0 if cv_set else 0.0
            else:
                score = (len(matched) / len(job_set)) * 100.0

            return score, matched, missing, extra

        def _calculate_experience_match(
            self,
            cv_years: Optional[float],
            job_years: Optional[int],
            job_range: Optional[Tuple[int, int]],
        ) -> float:
            if cv_years is None and job_years is None and job_range is None:
                return 50.0

            if cv_years is None:
                return 0.0

            if job_range:
                min_y, max_y = job_range
                if cv_years < min_y:
                    diff = min_y - cv_years
                    return max(0.0, 100.0 - diff * 20.0)
                if cv_years > max_y:
                    return 100.0
                return 100.0

            if job_years is not None:
                if cv_years >= job_years:
                    return 100.0
                diff = job_years - cv_years
                return max(0.0, 100.0 - diff * 20.0)

            return 50.0

        def _generate_recommendations(
            self,
            missing_skills: List[str],
            cv_data: Dict,
            job_data: Dict,
        ) -> List[str]:
            recs: List[str] = []

            if missing_skills:
                recs.append(
                    "Refuerza o adquiere las siguientes skills clave para este puesto: "
                    + ", ".join(missing_skills)
                )

            if cv_data.get("experience_years") is not None and job_data.get("experience_years"):
                if cv_data["experience_years"] < job_data["experience_years"]:
                    recs.append(
                        "Considera destacar proyectos o experiencias adicionales que muestren responsabilidad equivalente "
                        "aunque no sumen años formales."
                    )

            if not recs:
                recs.append("Tu perfil está bien alineado con la oferta. Ajusta el CV para resaltar los puntos fuertes.")

            return recs
    """
)

UI_GRADIO_PY = dedent(
    """\
    \"\"\"Interfaz mínima en Gradio para el MVP.\"\"\"

    import gradio as gr

    from src.cv_parser.cv_pipeline import CVPipeline
    from src.job_parser.job_pipeline import JobPipeline
    from src.matching.similarity import CVJobMatcher


    cv_pipeline = CVPipeline()
    job_pipeline = JobPipeline()
    matcher = CVJobMatcher()


    def analyze(cv_text: str, job_text: str) -> str:
        cv_data = cv_pipeline.process_text(cv_text) if hasattr(cv_pipeline, "process_text") else cv_pipeline.process("TODO")
        job_data = job_pipeline.process(job_text)

        result = matcher.calculate_match(
            cv_data=cv_data.__dict__,
            job_data=job_data.__dict__,
        )

        return (
            f"Score global: {result.global_score}\\n"
            f"Similitud semántica: {result.semantic_similarity}\\n"
            f"Match de skills: {result.skill_match_score}\\n"
            f"Match de experiencia: {result.experience_match_score}\\n\\n"
            f"Skills coincidentes: {', '.join(result.matched_skills) or '-'}\\n"
            f"Skills faltantes: {', '.join(result.missing_skills) or '-'}\\n"
            f"Recomendaciones:\\n- " + "\\n- ".join(result.recommendations)
        )


    def main():
        with gr.Blocks() as demo:
            gr.Markdown("# OPEN CAREER COACH - MVP")
            cv_input = gr.Textbox(label="Texto del CV", lines=10)
            job_input = gr.Textbox(label="Texto de la oferta", lines=10)
            output = gr.Textbox(label="Resultado", lines=15)

            btn = gr.Button("Analizar matching")
            btn.click(analyze, inputs=[cv_input, job_input], outputs=output)

        demo.launch()


    if __name__ == "__main__":
        main()
    """
)

README_MD = dedent(
    """\
    # OPEN CAREER COACH - MVP

    MVP para análisis de matching entre currículums y ofertas de empleo.

    ## Estructura

    - `src/config.py` — Configuración global.
    - `src/utils/file_loader.py` — Carga de archivos CV.
    - `src/cv_parser/cv_pipeline.py` — Parsing de CVs.
    - `src/job_parser/job_pipeline.py` — Parsing de ofertas.
    - `src/matching/similarity.py` — Motor de matching.
    - `src/ui/gradio_app.py` — Interfaz mínima en Gradio.

    ## Uso rápido

    ```bash
    pip install -r requirements.txt
    python -m src.ui.gradio_app
    ```
    """
)

# -------------------------
# Lógica del generador
# -------------------------

def create_directories(base_path: Path) -> None:
    for d in DIRECTORIES:
        (base_path / d).mkdir(parents=True, exist_ok=True)


def create_init_files(base_path: Path) -> None:
    for pkg in INIT_PACKAGES:
        init_file = base_path / pkg / "__init__.py"
        if not init_file.exists():
            init_file.write_text("", encoding="utf-8")


def write_file(path: Path, content: str, overwrite: bool = False) -> None:
    if path.exists() and not overwrite:
        return
    path.write_text(content.strip() + "\n", encoding="utf-8")


def generate_project(base_path: Path, overwrite: bool = False) -> None:
    create_directories(base_path)
    create_init_files(base_path)

    write_file(base_path / "requirements.txt", REQUIREMENTS_TXT, overwrite)
    write_file(base_path / "README.md", README_MD, overwrite)

    write_file(base_path / "src" / "config.py", CONFIG_PY, overwrite)
    write_file(base_path / "src" / "utils" / "file_loader.py", FILE_LOADER_PY, overwrite)
    write_file(base_path / "src" / "cv_parser" / "cv_pipeline.py", CV_PIPELINE_PY, overwrite)
    write_file(base_path / "src" / "job_parser" / "job_pipeline.py", JOB_PIPELINE_PY, overwrite)
    write_file(base_path / "src" / "matching" / "similarity.py", SIMILARITY_PY, overwrite)
    write_file(base_path / "src" / "ui" / "gradio_app.py", UI_GRADIO_PY, overwrite)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generador del MVP de OPEN CAREER COACH")
    parser.add_argument(
        "--path",
        type=str,
        default="./open-career-coach",
        help="Ruta base donde generar el proyecto",
    )
    parser.add_argument(
        "--overwrite",
        action="store_true",
        help="Sobrescribir archivos existentes",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    base_path = Path(args.path).resolve()
    os.makedirs(base_path, exist_ok=True)

    generate_project(base_path, overwrite=args.overwrite)

    print(f"MVP generado en: {base_path}")


if __name__ == "__main__":
    main()
