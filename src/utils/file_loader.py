"""Carga y extrae texto de archivos PDF, DOCX y TXT."""

from pathlib import Path
from typing import Dict


def load_text_from_file(filepath: str) -> Dict[str, str]:
    """Detecta la extensión del archivo y extrae el texto plano.

    Args:
        filepath: Ruta al archivo (PDF, DOCX o TXT)

    Returns:
        Dict con 'text', 'format', 'pages_estimated', 'characters'
    """
    path = Path(filepath)

    if not path.exists():
        raise FileNotFoundError(f"Archivo no encontrado: {filepath}")

    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf(filepath)
    elif suffix == ".docx":
        return _extract_docx(filepath)
    elif suffix == ".txt":
        return _extract_txt(filepath)
    else:
        raise ValueError(f"Formato no soportado: {suffix}. Use PDF, DOCX o TXT.")


def _extract_pdf(filepath: str) -> Dict[str, str]:
    import PyPDF2

    text_parts = []

    with open(filepath, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        total_pages = len(reader.pages)

        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

    full_text = "\n\n".join(text_parts)

    return {
        "text": full_text,
        "format": "PDF",
        "pages_estimated": total_pages,
        "characters": len(full_text),
    }


def _extract_docx(filepath: str) -> Dict[str, str]:
    from docx import Document

    doc = Document(filepath)

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                table_texts.append(row_text)

    full_text = "\n\n".join(paragraphs + table_texts)

    return {
        "text": full_text,
        "format": "DOCX",
        "pages_estimated": len(paragraphs) // 40 + 1,
        "characters": len(full_text),
    }


def _extract_txt(filepath: str) -> Dict[str, str]:
    with open(filepath, "r", encoding="utf-8") as f:
        text = f.read()

    if not text.strip():
        for encoding in ["latin-1", "iso-8859-1", "cp1252"]:
            try:
                with open(filepath, "r", encoding=encoding) as f:
                    text = f.read()
                if text.strip():
                    break
            except Exception:
                continue

    lines = text.splitlines()

    return {
        "text": text,
        "format": "TXT",
        "pages_estimated": len(lines) // 50 + 1,
        "characters": len(text),
    }
