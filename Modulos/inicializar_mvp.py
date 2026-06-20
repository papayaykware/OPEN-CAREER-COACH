# inicializar_mvp.py
# Generando el MVP funcional siguiendo la recomendación de DeepSeek
# con visión arquitectónica propia para OPEN-CAREER-COACH

import os

# Detecta automáticamente la carpeta 'Modulos' donde se ejecuta el script
base_path = os.path.dirname(os.path.abspath(__file__))

# ==========================================
# Fase 1: Crear estructura de directorios
# ==========================================
os.makedirs(f"{base_path}/src/utils", exist_ok=True)
os.makedirs(f"{base_path}/src/cv_parser", exist_ok=True)
os.makedirs(f"{base_path}/src/job_parser", exist_ok=True)
os.makedirs(f"{base_path}/src/matching", exist_ok=True)
os.makedirs(f"{base_path}/src/ui", exist_ok=True)
os.makedirs(f"{base_path}/data/sample_cvs", exist_ok=True)
os.makedirs(f"{base_path}/data/sample_jobs", exist_ok=True)

# Crear archivos __init__.py para empaquetado Python
modulos_python = ["src", "src/utils", "src/cv_parser", "src/job_parser", "src/matching", "src/ui"]
for d in modulos_python:
    init_file = f"{base_path}/{d}/__init__.py"
    if not os.path.exists(init_file):
        with open(init_file, "w") as f:
            f.write("")

print("🚀 Estructura de directorios MVP creada.")

# ==========================================
# Fase 2: Generar requirements.txt actualizado
# ==========================================
requirements = """# OPEN CAREER COACH - MVP Dependencies
# Core
gradio==4.44.0
sentence-transformers==3.0.1
PyPDF2==3.0.1
python-docx==1.1.2
chromadb==0.5.5
numpy==1.26.4

# NLP
spacy==3.7.5

# Utilities
python-dotenv==1.0.1
requests==2.32.3

# Testing
pytest==8.3.2
"""

with open(f"{base_path}/requirements.txt", "w") as f:
    f.write(requirements)

# ==========================================
# Fase 3: Generar config.py
# ==========================================
config_py = """# src/config.py
\"\"\"Configuración central del sistema MVP.\"\"\"

from dataclasses import dataclass


@dataclass
class Config:
    \"\"\"Configuración global.\"\"\"
    
    # Modelos
    EMBEDDING_MODEL: str = "sentence-transformers/all-MiniLM-L6-v2"
    
    # Matching
    SIMILARITY_THRESHOLD: float = 0.50
    
    # Paths
    DATA_DIR: str = "./data"
    SAMPLE_CVS_DIR: str = "./data/sample_cvs"
    SAMPLE_JOBS_DIR: str = "./data/sample_jobs"


config = Config()
"""

with open(f"{base_path}/src/config.py", "w") as f:
    f.write(config_py)

# ==========================================
# Fase 4: Generar utils/file_loader.py
# ==========================================
file_loader = '''# src/utils/file_loader.py
"""Carga y extrae texto de archivos PDF, DOCX y TXT."""

from pathlib import Path
from typing import Dict


def load_text_from_file(filepath: str) -> Dict[str, str]:
    """Detecta la extensión del archivo y extrae el texto plano.
    
    Args:
        filepath: Ruta al archivo (PDF, DOCX o TXT)
        
    Returns:
        Dict con 'text', 'format', 'pages_estimated'
        
    Raises:
        FileNotFoundError: Si el archivo no existe
        ValueError: Si el formato no es soportado
    """
    path = Path(filepath)
    
    if not path.exists():
        raise FileNotFoundError(f"Archivo no encontrado: {filepath}")
    
    suffix = path.suffix.lower()
    
    if suffix == ".pdf":
        return _extract_pdf(filepath)
    elif suffix == ".docx":
        return _extract_docx(filepath)
    elif suffix == ".txt":
        return _extract_txt(filepath)
    else:
        raise ValueError(f"Formato no soportado: {suffix}. Use PDF, DOCX o TXT.")


def _extract_pdf(filepath: str) -> Dict[str, str]:
    """Extrae texto de PDF usando PyPDF2."""
    import PyPDF2
    
    text_parts = []
    
    with open(filepath, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        total_pages = len(reader.pages)
        
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    
    full_text = "\\n\\n".join(text_parts)
    
    return {
        "text": full_text,
        "format": "PDF",
        "pages_estimated": total_pages,
        "characters": len(full_text),
    }


def _extract_docx(filepath: str) -> Dict[str, str]:
    """Extrae texto de DOCX usando python-docx."""
    from docx import Document
    
    doc = Document(filepath)
    
    # Extraer párrafos
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    # Extraer tablas
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                table_texts.append(row_text)
    
    full_text = "\\n\\n".join(paragraphs + table_texts)
    
    return {
        "text": full_text,
        "format": "DOCX",
        "pages_estimated": len(paragraphs) // 40 + 1,
        "characters": len(full_text),
    }


def _extract_txt(filepath: str) -> Dict[str, str]:
    """Extrae texto de archivo plano TXT."""
    with open(filepath, "r", encoding="utf-8") as f:
        text = f.read()
    
    # Fallback a otros encodings
    if not text.strip():
        for encoding in ["latin-1", "iso-8859-1", "cp1252"]:
            try:
                with open(filepath, "r", encoding=encoding) as f:
                    text = f.read()
                if text.strip():
                    break
            except Exception:
                continue
    
    lines = text.splitlines()
    
    return {
        "text": text,
        "format": "TXT",
        "pages_estimated": len(lines) // 50 + 1,
        "characters": len(text),
    }
'''

with open(f"{base_path}/src/utils/file_loader.py", "w") as f:
    f.write(file_loader)

# ==========================================
# Fase 5: Generar cv_parser/cv_pipeline.py
# ==========================================
cv_pipeline = '''# src/cv_parser/cv_pipeline.py
"""Pipeline de procesamiento de currículums para el MVP.

Extrae texto, identifica skills básicas y estructura la información.
"""

import re
from typing import Dict, List, Set, Optional
from dataclasses import dataclass, field


@dataclass
class CVData:
    """Datos estructurados extraídos de un CV."""
    raw_text: str = ""
    skills_technical: List[str] = field(default_factory=list)
    skills_soft: List[str] = field(default_factory=list)
    experience_years: Optional[float] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    linkedin: Optional[str] = None
    file_format: str = ""
    pages: int = 0


class CVPipeline:
    """Pipeline completo de procesamiento de CVs.
    
    Procesa archivos PDF, DOCX y TXT para extraer:
    - Texto completo
    - Skills técnicas (hard skills)
    - Skills blandas (soft skills)
    - Información de contacto
    - Años de experiencia estimados
    """
    
    # Diccionario de skills técnicas para detección
    TECH_SKILLS = {
        "python", "java", "javascript", "typescript", "c++", "c#", "go", "rust",
        "ruby", "php", "scala", "kotlin", "swift", "objective-c", "perl", "r",
        "matlab", "sas", "vba", "shell", "bash", "powershell",
        "html", "css", "sass", "less", "react", "vue", "vue.js", "angular",
        "svelte", "next.js", "nuxt", "django", "flask", "fastapi", "spring",
        "express", "nodejs", "node.js", "laravel", "rails", "asp.net",
        "sql", "mysql", "postgresql", "oracle", "mongodb", "redis",
        "elasticsearch", "cassandra", "dynamodb", "neo4j", "sqlite",
        "firebase", "supabase", "snowflake", "bigquery",
        "aws", "amazon web services", "azure", "gcp", "google cloud",
        "docker", "kubernetes", "openshift", "terraform", "ansible",
        "jenkins", "gitlab ci", "github actions", "circleci", "travis ci",
        "puppet", "chef", "vagrant", "nginx", "apache", "istio", "helm",
        "tensorflow", "pytorch", "keras", "scikit-learn", "xgboost",
        "lightgbm", "pandas", "numpy", "scipy", "matplotlib", "seaborn",
        "plotly", "tableau", "power bi", "looker", "spark", "hadoop",
        "kafka", "airflow", "dbt", "mlflow", "kubeflow",
        "react native", "flutter", "ionic", "cordova", "xamarin",
        "android", "ios", "pytest", "junit", "selenium", "cypress", "jest", "mocha",
        "cucumber", "postman", "jmeter", "k6", "git", "svn", "linux", "ubuntu", "windows", "macos",
        "rest api", "graphql", "grpc", "soap", "websockets",
        "microservices", "serverless", "lambda", "event-driven",
        "ci/cd", "devops", "sre", "platform engineering",
        "agile", "scrum", "kanban", "safe", "jira", "confluence",
        "blockchain", "solidity", "web3", "smart contracts",
        "figma", "sketch", "adobe xd", "invision",
        "wordpress", "drupal", "magento", "shopify",
    }
    
    # Diccionario de soft skills
    SOFT_SKILLS = {
        "liderazgo", "leadership", "comunicación", "communication",
        "trabajo en equipo", "teamwork", "team player",
        "resolución de problemas", "problem solving",
        "pensamiento crítico", "critical thinking",
        "adaptabilidad", "adaptability", "flexibility", "flexible",
        "creatividad", "creativity", "innovation", "innovative",
        "gestión del tiempo", "time management",
        "empatía", "empathy", "negociación", "negotiation",
        "presentación", "presentation skills",
        "proactividad", "proactive", "self-starter",
        "autonomía", "autonomy", "self-motivated",
        "atención al detalle", "attention to detail",
        "gestión de proyectos", "project management",
        "análisis", "analytical thinking",
        "orientación a resultados", "results-oriented",
        "orientación al cliente", "customer-oriented",
        "aprendizaje continuo", "continuous learning",
        "pensamiento estratégico", "strategic thinking",
        "toma de decisiones", "decision making",
        "gestión de conflictos", "conflict resolution",
        "mentoría", "mentoring", "coaching",
    }
    


